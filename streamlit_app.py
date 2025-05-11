import streamlit as st
import pandas as pd
import numpy as np
import os
import zipfile
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
import tempfile
import pdfkit
from PyPDF2 import PdfMerger
import base64
from jinja2 import Environment, FileSystemLoader
from pypdf import PdfReader, PdfWriter
import num2words
import platform
import traceback
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Temporary directory
TEMP_DIR = tempfile.mkdtemp()

# Configure wkhtmltopdf
if platform.system() == "Windows":
    wkhtmltopdf_path = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
else:
    config = pdfkit.configuration()

# Set up Jinja2 environment
env = Environment(loader=FileSystemLoader("templates"), cache_size=0)

# Helper functions
def number_to_words(number):
    try:
        return num2words(number, lang='en_IN')
    except:
        return str(number)

def set_cell_border(cell, **kwargs):
    """
    Set borders for a table cell in a Word document.
    Args:
        cell: The cell object from python-docx.
        kwargs: Dictionary with border settings for 'top', 'left', 'bottom', 'right'.
                Each border can have 'val', 'sz', 'color' attributes.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            edge_element = OxmlElement(f'w:{edge}')
            for key in edge_data:
                edge_element.set(qn(f'w:{key}'), str(edge_data[key]))
            # Remove any existing border for this edge to avoid duplicates
            existing_edge = tcBorders.find(qn(f'w:{edge}'))
            if existing_edge is not None:
                tcBorders.remove(existing_edge)
            tcBorders.append(edge_element)

def merge_pdfs(pdf_files, output_path):
    merger = PdfMerger()
    for pdf in pdf_files:
        if os.path.exists(pdf):
            merger.append(pdf)
    merger.write(output_path)
    merger.close()

import pandas as pd
import numpy as np

def process_bill(ws_wo, ws_bq, ws_extra, premium_percent, premium_type, amount_paid_last_bill, is_first_bill, user_inputs):
    try:
        # Initialize output structures
        data = {"items": []}  # For First Page table
        deviation_data = {"items": [], "summary": {}}  # For Deviation Statement
        header_data = {
            "deviation_headers": [
                "DEVIATION STATEMENT",
                f"Name of work: {user_inputs.get('work_name', '')}",
                f"Name of Contractor: {user_inputs.get('contractor_name', '')}",
                f"Agreement No.: {user_inputs.get('agreement_no', '')}"
            ],
            "tender_premium_bill": 0
        }

        # Read Work Order and Bill sheets
        ws_bill = ws_bq  # Using ws_bq as the bill sheet

        def is_header_or_invalid(value):
            """Check if a value indicates a header or invalid data."""
            if pd.isna(value):
                return True
            value_str = str(value).lower().strip()
            header_keywords = [
                "qty", "quantity", "rate", "amount", "sno", "serial", "unit",
                "description", "item", "total", "grand", "sub", "header"
            ]
            return any(keyword in value_str for keyword in header_keywords)

        def clean_numeric(value):
            """Clean and convert a value to float, return None if invalid."""
            if pd.isna(value) or value == "":
                return 0.0
            value_str = str(value).replace('%', '').strip()
            if is_header_or_invalid(value_str):
                return None
            try:
                return float(value_str)
            except ValueError:
                return None

        # Process Work Order items (start from row 21, 0-based index)
        wo_items = []
        for i in range(21, ws_wo.shape[0]):  # Changed from 20 to 21
            serial_no = str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else ""
            description = str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else ""
            unit = str(ws_wo.iloc[i, 2]) if pd.notnull(ws_wo.iloc[i, 2]) else ""
            
            # Clean and convert numeric values
            qty_wo = clean_numeric(ws_wo.iloc[i, 3])
            rate = clean_numeric(ws_wo.iloc[i, 4])
            amount = clean_numeric(ws_wo.iloc[i, 5])
            bsr = str(ws_wo.iloc[i, 6]) if pd.notnull(ws_wo.iloc[i, 6]) else ""

            # Skip invalid rows
            if not serial_no and not description and not unit:
                continue
            if qty_wo is None or rate is None or amount is None:
                st.warning(f"Skipping Work Order row {i+2}: Invalid numeric value (qty={ws_wo.iloc[i, 3]}, rate={ws_wo.iloc[i, 4]}, amount={ws_wo.iloc[i, 5]})")
                continue

            item = {
                "serial_no": serial_no,
                "description": description,
                "unit": unit,
                "qty_wo": qty_wo,
                "rate": rate,
                "amount": amount,
                "bsr": bsr
            }
            wo_items.append(item)

        # Process Bill items (start from row 21, 0-based index)
        bill_items = []
        for i in range(21, ws_bill.shape[0]):  # Changed from 20 to 21
            serial_no = str(ws_bill.iloc[i, 0]) if pd.notnull(ws_bill.iloc[i, 0]) else ""
            description = str(ws_bill.iloc[i, 1]) if pd.notnull(ws_bill.iloc[i, 1]) else ""
            unit = str(ws_bill.iloc[i, 2]) if pd.notnull(ws_bill.iloc[i, 2]) else ""
            
            # Clean and convert numeric values
            qty_bill = clean_numeric(ws_bill.iloc[i, 3])
            rate = clean_numeric(ws_bill.iloc[i, 4])
            amount = clean_numeric(ws_bill.iloc[i, 5])
            bsr = str(ws_bill.iloc[i, 6]) if pd.notnull(ws_bill.iloc[i, 6]) else ""

            # Skip invalid rows
            if not serial_no and not description and not unit:
                continue
            if qty_bill is None or rate is None or amount is None:
                st.warning(f"Skipping Bill Quantity row {i+2}: Invalid numeric value (qty={ws_bill.iloc[i, 3]}, rate={ws_bill.iloc[i, 4]}, amount={ws_bill.iloc[i, 5]})")
                continue

            item = {
                "serial_no": serial_no,
                "description": description,
                "unit": unit,
                "qty_bill": qty_bill,
                "rate": rate,
                "amount": amount,
                "bsr": bsr
            }
            bill_items.append(item)

        # Create data for First Page table
        for wo_item in wo_items:
            bill_item = next((bi for bi in bill_items if bi["bsr"] == wo_item["bsr"]), None)
            qty_bill = bill_item["qty_bill"] if bill_item else 0
            amount_bill = bill_item["amount"] if bill_item else 0
            item = {
                "unit": wo_item["unit"],
                "qty_since_last": qty_bill,  # Assuming first bill
                "qty_upto_date": qty_bill,
                "serial_no": wo_item["serial_no"],
                "description": wo_item["description"],
                "rate": wo_item["rate"],
                "amount_upto_date": amount_bill,
                "amount_since_prev": amount_bill,
                "remarks": ""
            }
            data["items"].append(item)

        # Create deviation_data for Deviation Statement table
        for wo_item in wo_items:
            bill_item = next((bi for bi in bill_items if bi["bsr"] == wo_item["bsr"]), None)
            qty_wo = wo_item["qty_wo"]
            rate = wo_item["rate"]
            amt_wo = wo_item["amount"]
            qty_bill = bill_item["qty_bill"] if bill_item else 0
            amt_bill = bill_item["amount"] if bill_item else 0
            excess_qty = max(0, qty_bill - qty_wo)
            excess_amt = excess_qty * rate
            saving_qty = max(0, qty_wo - qty_bill)
            saving_amt = saving_qty * rate
            item = {
                "serial_no": wo_item["serial_no"],
                "description": wo_item["description"],
                "unit": wo_item["unit"],
                "qty_wo": qty_wo,
                "rate": rate,
                "amt_wo": amt_wo,
                "qty_bill": qty_bill,
                "amt_bill": amt_bill,
                "excess_qty": excess_qty,
                "excess_amt": excess_amt,
                "saving_qty": saving_qty,
                "saving_amt": saving_amt,
                "remark": ""
            }
            deviation_data["items"].append(item)

        # Handle Extra Items
        if ws_extra is not None:
            for i in range(ws_extra.shape[0]):
                if pd.isna(ws_extra.iloc[i, 0]) or ws_extra.iloc[i, 0] == "":
                    continue
                serial_no = str(ws_extra.iloc[i, 0]) if pd.notnull(ws_extra.iloc[i, 0]) else ""
                description = str(ws_extra.iloc[i, 2]) if pd.notnull(ws_extra.iloc[i, 2]) else ""
                unit = str(ws_extra.iloc[i, 3]) if pd.notnull(ws_extra.iloc[i, 3]) else ""
                
                qty_bill = clean_numeric(ws_extra.iloc[i, 4])
                rate = clean_numeric(ws_extra.iloc[i, 5])
                amt_bill = clean_numeric(ws_extra.iloc[i, 6])
                
                if qty_bill is None or rate is None or amt_bill is None:
                    st.warning(f"Skipping Extra Items row {i+2}: Invalid numeric value (qty={ws_extra.iloc[i, 4]}, rate={ws_extra.iloc[i, 5]}, amount={ws_extra.iloc[i, 6]})")
                    continue
                
                item = {
                    "serial_no": serial_no,
                    "description": description,
                    "unit": unit,
                    "qty_wo": 0,  # Extra items not in Work Order
                    "rate": rate,
                    "amt_wo": 0,
                    "qty_bill": qty_bill,
                    "amt_bill": amt_bill,
                    "excess_qty": qty_bill,
                    "excess_amt": amt_bill,
                    "saving_qty": 0,
                    "saving_amt": 0,
                    "remark": ""
                }
                deviation_data["items"].append(item)

        return data, deviation_data, header_data

    except Exception as e:
        st.error(f"Error processing bill: {str(e)}")
        raise

def generate_bill_notes(payable_amount, work_order_amount, extra_item_amount):
    percentage_work_done = float(payable_amount / work_order_amount * 100) if work_order_amount > 0 else 0
    serial_number = 1
    note = []
    note.append(f"{serial_number}. The work has been completed {percentage_work_done:.2f}% of the Work Order Amount.")
    serial_number += 1
    if percentage_work_done < 90:
        note.append(f"{serial_number}. The execution of work at final stage is less than 90%...")
        serial_number += 1
    elif percentage_work_done > 100 and percentage_work_done <= 105:
        note.append(f"{serial_number}. Requisite Deviation Statement is enclosed...")
        serial_number += 1
    elif percentage_work_done > 105:
        note.append(f"{serial_number}. Requisite Deviation Statement is enclosed...")
        serial_number += 1
    note.append(f"{serial_number}. Quality Control (QC) test reports attached.")
    serial_number += 1
    if extra_item_amount > 0:
        extra_item_percentage = float(extra_item_amount / work_order_amount * 100) if work_order_amount > 0 else 0
        if extra_item_percentage > 5:
            note.append(f"{serial_number}. The amount of Extra items is Rs. {extra_item_amount}...")
        else:
            note.append(f"{serial_number}. The amount of Extra items is Rs. {extra_item_amount}...")
        serial_number += 1
    note.append(f"{serial_number}. Please peruse above details for necessary decision-making.")
    note.append("")
    note.append("                                Premlata Jain")
    note.append("                               AAO- As Auditor")
    return {
        "notes": note,
        "work_order_amount": work_order_amount,
        "totals": {
            "grand_total": payable_amount,
            "premium": {"percent": 0.0, "type": "above", "amount": 0},
            "payable": payable_amount,
            "extra_items_sum": extra_item_amount
        }
    }

def generate_pdf(sheet_name, data, orientation, output_path):
    try:
        required_fields = {
            "First Page": ["header", "items", "totals"],
            "Last Page": ["payable_amount", "amount_words"],
            "Deviation Statement": ["items", "summary", "header"],
            "Extra Items": ["items"],
            "Note Sheet": [
                "agreement_no", "name_of_work", "name_of_firm",
                "date_commencement", "date_completion", "actual_completion",
                "work_order_amount", "extra_item_amount", "notes", "totals"
            ],
            "Certificate III": [
                "payable_amount", "total_123", "balance_4_minus_5",
                "amount_paid_last_bill", "payment_now", "by_cheque",
                "cheque_amount_words", "certificate_items",
                "total_recovery", "totals"
            ]
        }

        required = required_fields.get(sheet_name, [])
        for field in required:
            if field not in data:
                raise ValueError(f"Missing required field for {sheet_name}: {field}")

        if "totals" in required and "totals" in data:
            required_totals = {
                "First Page": ["grand_total", "premium", "payable"],
                "Certificate III": ["grand_total", "payable_amount", "extra_items_sum", "total_123"]
            }
            required = required_totals.get(sheet_name, [])
            for field in required:
                if field not in data["totals"]:
                    raise ValueError(f"Missing required totals field for {sheet_name}: {field}")

        template = env.get_template(f"{sheet_name.lower().replace(' ', '_')}.html")
        html_content = template.render(data=data)

        debug_html_path = os.path.join(TEMP_DIR, f"{sheet_name.replace(' ', '_')}_debug.html")
        with open(debug_html_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        options = {
            "page-size": "A4",
            "orientation": orientation,
            "margin-top": "0.25in",
            "margin-bottom": "0.25in",
            "margin-left": "0.25in",
            "margin-right": "0.5in",
            "encoding": "UTF-8",
            "quiet": "",
            "no-outline": None,
            "enable-local-file-access": None,
            "disable-smart-shrinking": None,
            "dpi": 300,
            "javascript-delay": "1000",
            "no-stop-slow-scripts": None,
            "load-error-handling": "ignore"
        }

        if sheet_name == "Note Sheet":
            options["margin-bottom"] = "0.6in"
        elif sheet_name == "Deviation Statement":
            options["margin-bottom"] = "0.25in"

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        pdfkit.from_string(html_content, output_path, configuration=config, options=options)

        if not os.path.exists(output_path):
            raise Exception(f"PDF file was not created at {output_path}")

        return True

    except Exception as e:
        st.error(f"Error generating PDF for {sheet_name}: {str(e)}")
        st.write(traceback.format_exc())
        raise

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
import traceback

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
import traceback

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
import traceback

def create_word_doc(data, deviation_data, header_data, sheet_name=""):
    try:
        # Validate input types
        if not isinstance(data, dict):
            raise TypeError(f"Expected 'data' to be a dictionary, got {type(data).__name__}: {data}")
        if "items" not in data:
            raise KeyError("'data' dictionary missing 'items' key")
        if not isinstance(deviation_data, dict):
            raise TypeError(f"Expected 'deviation_data' to be a dictionary, got {type(deviation_data).__name__}")
        if "items" not in deviation_data:
            raise KeyError("'deviation_data' dictionary missing 'items' key")

        doc = Document()
        
        # First Page table
        section = doc.sections[0]
        section.left_margin = Inches(0.25)
        section.right_margin = Inches(0.5)
        table = doc.add_table(rows=len(data["items"]) + 4, cols=9)
        table.style = "Table Grid"
        column_widths = [0.4, 0.55, 0.55, 0.38, 2.5, 0.52, 0.77, 0.6, 0.47]
        for i, width in enumerate(column_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)
        table.cell(0, 0).text = "Unit"
        table.cell(0, 1).text = "Quantity executed (or supplied) since last certificate"
        table.cell(0, 2).text = "Quantity executed (or supplied) upto date as per MB"
        table.cell(0, 3).text = "S. No."
        table.cell(0, 4).text = 'Item of Work supplies (Grouped under "sub-head" and "sub work" of estimate)'
        table.cell(0, 5).text = "Rate"
        table.cell(0, 6).text = "Upto date Amount"
        table.cell(0, 7).text = "Amount Since previous bill (Total for each sub-head)"
        table.cell(0, 8).text = "Remarks"
        # Populate First Page data (omitted for brevity)
        
        # Deviation Statement table
        section = doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.5512)
        section.right_margin = Inches(0.5512)
        
        table = doc.add_table(rows=len(deviation_data["items"]) + 5, cols=13)
        table.style = "Table Grid"
        table.autofit = False
        
        column_widths = [0.353, 6.352, 0.353, 0.353, 0.353, 0.353, 0.353, 0.353, 0.353, 0.353, 0.353, 0.353, 0.795]
        for i, width in enumerate(column_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)
        
        table.cell(0, 0).text = "ITEM No."
        table.cell(0, 1).text = "Description"
        table.cell(0, 2).text = "Unit"
        table.cell(0, 3).text = "Qty as per Work Order"
        table.cell(0, 4).text = "Rate"
        table.cell(0, 5).text = "Amt as per Work Order Rs."
        table.cell(0, 6).text = "Qty Executed"
        table.cell(0, 7).text = "Amt as per Executed Rs."
        table.cell(0, 8).text = "Excess Qty"
        table.cell(0, 9).text = "Excess Amt Rs."
        table.cell(0, 10).text = "Saving Qty"
        table.cell(0, 11).text = "Saving Amt Rs."
        table.cell(0, 12).text = "REMARKS/REASON."
        
        for i, item in enumerate(deviation_data["items"], 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(item.get("serial_no", ""))
            row_cells[1].text = item.get("description", "")
            row_cells[2].text = item.get("unit", "")
            row_cells[3].text = str(item.get("qty_wo", "")) if item.get("unit", "").strip() else ""
            row_cells[4].text = str(item.get("rate", "")) if item.get("unit", "").strip() else ""
            row_cells[5].text = str(item.get("amt_wo", "")) if item.get("unit", "").strip() and item.get("rate", "").strip() else ""
            row_cells[6].text = str(item.get("qty_bill", "")) if item.get("unit", "").strip() else ""
            row_cells[7].text = str(item.get("amt_bill", "")) if item.get("unit", "").strip() and item.get("rate", "").strip() else ""
            row_cells[8].text = str(item.get("excess_qty", "")) if item.get("unit", "").strip() else ""
            row_cells[9].text = str(item.get("excess_amt", "")) if item.get("unit", "").strip() and item.get("rate", "").strip() else ""
            row_cells[10].text = str(item.get("saving_qty", "")) if item.get("unit", "").strip() else ""
            row_cells[11].text = str(item.get("saving_amt", "")) if item.get("unit", "").strip() and item.get("rate", "").strip() else ""
            row_cells[12].text = item.get("remark", "")
        
        summary_row = table.rows[-5].cells
        summary_row[1].text = "Grand Total Rs."
        summary_row[5].text = str(deviation_data["summary"].get("work_order_total", ""))
        summary_row[7].text = str(deviation_data["summary"].get("executed_total", ""))
        summary_row[9].text = str(deviation_data["summary"].get("overall_excess", ""))
        summary_row[11].text = str(deviation_data["summary"].get("overall_saving", ""))
        
        premium_row = table.rows[-4].cells
        premium_percent = deviation_data["summary"].get("premium", {}).get("percent", 0) * 100
        premium_row[1].text = f"Add Tender Premium ({premium_percent:.2f}%)"
        premium_row[5].text = str(deviation_data["summary"].get("tender_premium_f", ""))
        premium_row[7].text = str(deviation_data["summary"].get("tender_premium_h", ""))
        premium_row[9].text = str(deviation_data["summary"].get("tender_premium_j", ""))
        premium_row[11].text = str(deviation_data["summary"].get("tender_premium_l", ""))
        
        grand_total_row = table.rows[-3].cells
        grand_total_row[1].text = "Grand Total including Tender Premium Rs."
        grand_total_row[5].text = str(deviation_data["summary"].get("grand_total_f", ""))
        grand_total_row[7].text = str(deviation_data["summary"].get("grand_total_h", ""))
        grand_total_row[9].text = str(deviation_data["summary"].get("grand_total_j", ""))
        grand_total_row[11].text = str(deviation_data["summary"].get("grand_total_l", ""))
        
        net_diff_row = table.rows[-2].cells
        net_diff = deviation_data["summary"].get("net_difference", 0)
        net_diff_row[1].text = "Overall Excess With Respect to the Work Order Amount Rs." if net_diff > 0 else "Overall Saving With Respect to the Work Order Amount Rs."
        net_diff_row[7].text = str(abs(net_diff))
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(8)
        
        doc.save("output.docx")
    
    except Exception as e:
        st.error(f"Error generating Word document for {sheet_name if sheet_name else 'Deviation Statement'}: {str(e)}")
        st.write(traceback.format_exc())
        raise

def main():
    st.markdown("""
    <style>
        .stButton > button {
            background-color: #4CAF50;
            color: white;
            padding: 10px 20px;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            font-size: 16px;
        }
        .stButton > button:hover {
            background-color: #45a049;
        }
        .stFileUploader > div > div {
            background-color: #f5f5f5;
            padding: 10px;
            border-radius: 4px;
            border: 1px solid #ddd;
        }
        .stFileUploader > div > div:hover {
            background-color: #e0e0e0;
        }
        .stFormSubmitButton > button {
            background-color: #2196F3 !important;
            color: white !important;
        }
        .stFormSubmitButton > button:hover {
            background-color: #1976D2 !important;
        }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style='text-align: center; padding: 20px; background-color: #f9f9f9; border-radius: 8px;'>
        <h3>Generate Contractor Bills with Ease</h3>
        <p>Upload your Excel files and generate professional contractor bills in seconds.</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    ### Instructions:
    1. Fill in the required details in the sidebar
    2. Upload an Excel file containing three sheets:
       - Work Order (ws_wo)
       - Bill Quantity (ws_bq)
       - Extra Items (ws_extra)
    """, unsafe_allow_html=True)

    uploaded_file = st.file_uploader("Upload Excel File", type=["xlsx", "xls"])

    fixed_header = "FOR CONTRACTORS & SUPPLIERS ONLY FOR PAYMENT FOR WORK OR SUPPLIES ACTUALLY MEASURED WORK ORDER"
    bill_serial = st.sidebar.text_input("Serial No. of this bill", "First & Final Bill")
    start_date = st.sidebar.date_input("St. date of Start")
    completion_date = st.sidebar.date_input("St. date of completion")
    actual_completion_date = st.sidebar.date_input("Date of actual completion of work")
    work_order_amount = st.sidebar.number_input("WORK ORDER AMOUNT RS.", min_value=0.0)
    premium_type = st.sidebar.selectbox("Premium Type", ["Above", "Below"])
    premium_percent = st.sidebar.number_input("Premium Percentage", min_value=0.0, max_value=100.0, value=0.0, step=0.1)
    amount_paid_last_bill = st.sidebar.number_input("Amount Paid vide Last bill", min_value=0.0)
    cash_voucher_no = st.sidebar.text_input("Cash Book Voucher No.", "")
    cash_voucher_date = st.sidebar.date_input("Cash Book Voucher Date", value=None)
    contractor_name = st.sidebar.text_input("Name of Contractor or supplier", "")
    work_description = st.sidebar.text_area("Name of Work", "")
    last_bill_no = st.sidebar.text_input("No. and date of the last bill", "Not Applicable")
    work_order_ref = st.sidebar.text_input("Reference to work order or Agreement", "")
    agreement_no = st.sidebar.text_input("Agreement No.", "")
    written_order_date = st.sidebar.date_input("Date of written order to commence work", value=None)
    is_first_bill = st.sidebar.checkbox("Is First Bill")

    if st.button("Generate Bill"):
        if uploaded_file is not None:
            try:
                if not all([bill_serial, start_date, completion_date, actual_completion_date, work_order_amount]):
                    st.error("Please fill all mandatory fields")
                    return

                with pd.ExcelFile(uploaded_file) as xls:
                    ws_wo = pd.read_excel(xls, "Work Order", header=None)
                    ws_bq = pd.read_excel(xls, "Bill Quantity", header=None)
                    ws_extra = pd.read_excel(xls, "Extra Items", header=None)

                first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data, certificate_iii_data = process_bill(
                    ws_wo,
                    ws_bq,
                    ws_extra,
                    premium_percent,
                    premium_type,
                    amount_paid_last_bill,
                    is_first_bill,
                    user_inputs={
                        "fixed_header": fixed_header,
                        "bill_serial": bill_serial,
                        "start_date": start_date.strftime("%d-%m-%Y") if start_date else "",
                        "completion_date": completion_date.strftime("%d-%m-%Y") if completion_date else "",
                        "actual_completion_date": actual_completion_date.strftime("%d-%m-%Y") if actual_completion_date else "",
                        "work_order_amount": work_order_amount,
                        "premium_percent": premium_percent,
                        "premium_type": premium_type,
                        "amount_paid_last_bill": amount_paid_last_bill,
                        "cash_voucher_no": cash_voucher_no,
                        "cash_voucher_date": cash_voucher_date.strftime("%d-%m-%Y") if cash_voucher_date else "",
                        "contractor_name": contractor_name,
                        "work_description": work_description,
                        "last_bill_no": last_bill_no,
                        "work_order_ref": work_order_ref,
                        "agreement_no": agreement_no,
                        "written_order_date": written_order_date.strftime("%d-%m-%Y") if written_order_date else "",
                        "is_first_bill": is_first_bill,
                        "measurement_officer": "Measurement Officer Name",
                        "measurement_date": "30/04/2025",
                        "measurement_book_page": "123",
                        "measurement_book_no": "MB-001",
                        "officer_name": "Officer Name",
                        "officer_designation": "Designation",
                        "authorising_officer_name": "Authorising Officer Name",
                        "authorising_officer_designation": "Designation"
                    }
                )

                pdf_files = []
                for sheet_name, data, orientation, template_name in [
                    ("First Page", first_page_data, "portrait", "first_page"),
                    ("Last Page", last_page_data, "portrait", "last_page"),
                    ("Extra Items", extra_items_data, "portrait", "extra_items"),
                    ("Deviation Statement", deviation_data, "landscape", "deviation_statement"),
                    ("Note Sheet", note_sheet_data, "portrait", "note_sheet"),
                    ("Certificate III", certificate_iii_data, "portrait", "certificate_iii")
                ]:
                    pdf_path = os.path.join(TEMP_DIR, f"{sheet_name.replace(' ', '_')}.pdf")
                    if generate_pdf(template_name, data, orientation, pdf_path):
                        pdf_files.append(pdf_path)

                current_date = datetime.now().strftime("%Y%m%d")
                pdf_output = os.path.join(TEMP_DIR, f"BILL_AND_DEVIATION_{current_date}.pdf")
                merge_pdfs(pdf_files, pdf_output)

                word_files = []
                for sheet_name, data in [
                    ("First Page", first_page_data),
                    ("Last Page", last_page_data),
                    ("Extra Items", extra_items_data),
                    ("Deviation Statement", deviation_data),
                    ("Note Sheet", note_sheet_data),
                    ("Certificate III", certificate_iii_data)
                ]:
                    doc_path = os.path.join(TEMP_DIR, f"{sheet_name.replace(' ', '_')}.docx")
                    if create_word_doc(sheet_name, data, doc_path, "landscape" if sheet_name == "Deviation Statement" else "portrait"):
                        word_files.append(doc_path)

                zip_path = os.path.join(TEMP_DIR, "output.zip")
                with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
                    zipf.write(pdf_output, os.path.basename(pdf_output))
                    for word_file in word_files:
                        zipf.write(word_file, os.path.basename(word_file))

                with open(zip_path, "rb") as f:
                    bytes_data = f.read()
                st.download_button(
                    label="Download Output Files",
                    data=bytes_data,
                    file_name="bill_output.zip",
                    mime="application/zip"
                )
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")
                st.stop()

if __name__ == "__main__":
    main()